"""
Генератор листинга кода для дипломной работы «Farme».
Читает исходные файлы напрямую и создаёт LISTING_FARME.docx на рабочем столе.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = os.path.dirname(os.path.abspath(__file__))
import ctypes.wintypes, ctypes
CSIDL_DESKTOP = 0
buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
ctypes.windll.shell32.SHGetFolderPathW(0, CSIDL_DESKTOP, 0, 0, buf)
DESKTOP = buf.value or os.path.join(os.path.expanduser("~"), "Desktop")
OUT_PATH = os.path.join(DESKTOP, "LISTING_FARME.docx")
SRC_BASE = os.path.join(PROJECT, "app", "src", "main", "java", "com", "example", "farme")

# ─── Структура листинга ──────────────────────────────────────────────────────
# (section, section_title, filename, relative_path_from_farme, description)
# section=None → продолжение предыдущего раздела

SECTIONS = [
    # ══ РАЗДЕЛ 1: Базовые классы ══
    ("1", "Базовые классы приложения",
     "FarmeApp.java", "FarmeApp.java",
     "Класс Application — инициализация темы, каналов уведомлений, FCM-токена при старте"),

    (None, None,
     "BaseActivity.java", "BaseActivity.java",
     "Абстрактный базовый класс для всех Activity — применяет выбранную локаль (ru/en/ky)"),

    (None, None,
     "LocaleHelper.java", "LocaleHelper.java",
     "Вспомогательный класс смены языка — оборачивает Context с нужной локалью"),

    (None, None,
     "SplashActivity.java", "SplashActivity.java",
     "Заставка с анимацией — маршрутизация: онбординг / авторизация / главный экран"),

    (None, None,
     "NotificationHelper.java", "NotificationHelper.java",
     "Создание notification-каналов (channel_chat, channel_system) для Android 8+"),

    (None, None,
     "FarmeFcmService.java", "FarmeFcmService.java",
     "Firebase Cloud Messaging сервис — получение push-уведомлений и сохранение токена"),

    (None, None,
     "PhoneVerificationHelper.java", "PhoneVerificationHelper.java",
     "Заглушка проверки поддержки Phone Auth — всегда направляет на ручной ввод SMS OTP"),

    # ══ РАЗДЕЛ 2: Аутентификация ══
    ("2", "Модуль аутентификации",
     "auth/AuthActivity.java", os.path.join("auth", "AuthActivity.java"),
     "Экран входа/регистрации — вкладки логин/регистрация, методы телефон/email"),

    (None, None,
     "auth/OtpActivity.java", os.path.join("auth", "OtpActivity.java"),
     "Ввод SMS OTP — 6 отдельных полей, CountDownTimer 60 с, автозаполнение из SMS"),

    (None, None,
     "auth/CompleteProfileActivity.java", os.path.join("auth", "CompleteProfileActivity.java"),
     "Заполнение профиля после регистрации — имя, фамилия, регион, пароль"),

    (None, None,
     "auth/ForgotPasswordActivity.java", os.path.join("auth", "ForgotPasswordActivity.java"),
     "Восстановление пароля — отправка письма сброса через Firebase Auth"),

    # ══ РАЗДЕЛ 3: Основные экраны ══
    ("3", "Основные экраны приложения",
     "MainActivity.java", "MainActivity.java",
     "Главная Activity — нижняя навигация (Home/Chats/FAB/Map/Profile), бейдж непрочитанных"),

    (None, None,
     "OnboardingActivity.java", "OnboardingActivity.java",
     "Онбординг из 4 слайдов на ViewPager2 с точечными индикаторами"),

    (None, None,
     "CreateListingActivity.java", "CreateListingActivity.java",
     "Мастер создания объявления (3 шага): фото → категория+местоположение → превью+публикация"),

    (None, None,
     "ListingDetailActivity.java", "ListingDetailActivity.java",
     "Детальная карточка объявления — галерея, продавец, кнопки звонка/чата/архива, карта"),

    (None, None,
     "EditListingActivity.java", "EditListingActivity.java",
     "Редактирование объявления — изменение полей, сброс статуса модерации"),

    (None, None,
     "ChatActivity.java", "ChatActivity.java",
     "Экран чата — real-time сообщения Firebase, свайп для удаления, счётчик непрочитанных"),

    (None, None,
     "SellerProfileActivity.java", "SellerProfileActivity.java",
     "Профиль продавца — объявления/отзывы, блокировка, жалоба, звонок/чат"),

    (None, None,
     "ProfileActivity.java", "ProfileActivity.java",
     "Личный профиль пользователя — статистика, меню, аватар, счётчики"),

    (None, None,
     "EditProfileActivity.java", "EditProfileActivity.java",
     "Редактирование профиля — имя, регион, загрузка аватара (Base64)"),

    (None, None,
     "MyListingsActivity.java", "MyListingsActivity.java",
     "Мои объявления — вкладки Активные/На модерации/Архив"),

    (None, None,
     "FavoritesActivity.java", "FavoritesActivity.java",
     "Избранные объявления — загрузка из Firebase, удаление из избранного"),

    (None, None,
     "SalesActivity.java", "SalesActivity.java",
     "История продаж — загрузка проданных объявлений из узла sales/{uid}"),

    (None, None,
     "ReviewsActivity.java", "ReviewsActivity.java",
     "Отзывы о продавце — средний рейтинг, список отзывов из reviews/{uid}"),

    (None, None,
     "NotificationsActivity.java", "NotificationsActivity.java",
     "Уведомления — real-time, свайп для удаления, отметка прочитанными"),

    (None, None,
     "SettingsActivity.java", "SettingsActivity.java",
     "Настройки — тема, язык, регион, валюта, push-уведомления, блоклист, выход"),

    (None, None,
     "SupportActivity.java", "SupportActivity.java",
     "Поддержка — FAQ-аккордеон, отправка обращений, просмотр тикетов"),

    (None, None,
     "TicketChatActivity.java", "TicketChatActivity.java",
     "Чат обращения в поддержку — пузырьковый UI, ответы администратора"),

    (None, None,
     "AdminSupportActivity.java", "AdminSupportActivity.java",
     "Панель администратора для тикетов — вкладки открытые/отвеченные, диалог ответа"),

    (None, None,
     "BlocklistActivity.java", "BlocklistActivity.java",
     "Список заблокированных пользователей — просмотр и разблокировка"),

    # ══ РАЗДЕЛ 4: Фрагменты ══
    ("4", "Фрагменты навигации",
     "fragments/HomeFragment.java", os.path.join("fragments", "HomeFragment.java"),
     "Главная лента — Firebase-загрузка объявлений, многокритериальная фильтрация, история"),

    (None, None,
     "fragments/MapFragment.java", os.path.join("fragments", "MapFragment.java"),
     "Карта — кастомные маркеры с ценой на Canvas, фильтр категорий, горизонтальный список"),

    (None, None,
     "fragments/ChatsFragment.java", os.path.join("fragments", "ChatsFragment.java"),
     "Список чатов — загрузка chatIds, превью, поиск, свайп для удаления, бейдж"),

    (None, None,
     "fragments/ProfileFragment.java", os.path.join("fragments", "ProfileFragment.java"),
     "Фрагмент профиля — аватар, рейтинг, статистика, меню, SwipeRefreshLayout"),

    (None, None,
     "fragments/FavoritesFragment.java", os.path.join("fragments", "FavoritesFragment.java"),
     "Фрагмент избранного — загрузка из Firebase, очистить всё, переход к ленте"),

    (None, None,
     "fragments/SearchFragment.java", os.path.join("fragments", "SearchFragment.java"),
     "Поиск объявлений — TextWatcher, фильтры по категории и региону через Spinner"),

    (None, None,
     "fragments/FilterBottomSheet.java", os.path.join("fragments", "FilterBottomSheet.java"),
     "Панель фильтров (BottomSheet) — категории, подкатегории, регионы, цена, рейтинг, переключатели"),

    # ══ РАЗДЕЛ 5: Модели данных ══
    ("5", "Модели данных",
     "model/Listing.java", os.path.join("model", "Listing.java"),
     "POJO объявления с вложенным классом Passport (данные животного, ветсвидетельство)"),

    (None, None,
     "model/User.java", os.path.join("model", "User.java"),
     "POJO пользователя — uid, name, phone, email, region, role, rating, avatar, banned"),

    (None, None,
     "model/Message.java", os.path.join("model", "Message.java"),
     "POJO сообщения чата — id, senderId, text, createdAt, read, type (text/image)"),

    # ══ РАЗДЕЛ 6: Адаптеры ══
    ("6", "Адаптеры RecyclerView",
     "adapter/ListingAdapter.java", os.path.join("adapter", "ListingAdapter.java"),
     "Адаптер карточек объявлений — Base64-фото через Glide, избранное, кнопка чата"),

    (None, None,
     "adapter/MessageAdapter.java", os.path.join("adapter", "MessageAdapter.java"),
     "Адаптер сообщений чата — два типа ViewHolder (своё/чужое), время"),

    (None, None,
     "adapter/PhotoAdapter.java", os.path.join("adapter", "PhotoAdapter.java"),
     "Адаптер предпросмотра фотографий при создании объявления — Uri, кнопка удаления"),

    (None, None,
     "adapter/PhotoPagerAdapter.java", os.path.join("adapter", "PhotoPagerAdapter.java"),
     "Адаптер галереи фото в деталях объявления — ViewPager2, Base64-декодирование"),

    (None, None,
     "adapter/FeaturedListingAdapter.java", os.path.join("adapter", "FeaturedListingAdapter.java"),
     "Адаптер рекомендованных объявлений — горизонтальная карусель на главной странице"),

    # ══ РАЗДЕЛ 7: Утилиты и диалоги ══
    ("7", "Утилиты и вспомогательные классы",
     "utils/Validator.java", os.path.join("utils", "Validator.java"),
     "Валидатор — regex для телефона +996, email, пароль, имя, цена; форматирование"),

    (None, None,
     "utils/HistoryManager.java", os.path.join("utils", "HistoryManager.java"),
     "История просмотров — SharedPreferences, до 20 ID объявлений на пользователя"),

    (None, None,
     "ReviewDialogHelper.java", "ReviewDialogHelper.java",
     "Диалог оставления отзыва — 5 звёзд, текст, проверка дубликата, пересчёт рейтинга"),

    (None, None,
     "RejectDialogHelper.java", "RejectDialogHelper.java",
     "Диалог отклонения объявления администратором — выбор причины, уведомление продавца"),
]

# ─── helpers ────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"Раздел {number}. {title}")
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_file_header(doc, listing_num, filename, path, description):
    doc.add_paragraph()
    sep = doc.add_paragraph("─" * 85)
    sep.runs[0].font.name  = "Courier New"
    sep.runs[0].font.size  = Pt(8)
    sep.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run(f"Листинг {listing_num} — {filename}")
    run.font.name  = "Times New Roman"
    run.font.bold  = True
    run.font.size  = Pt(11)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(0)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Путь: {path}")
    r2.font.name  = "Courier New"
    r2.font.size  = Pt(9)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)

    p3 = doc.add_paragraph()
    r3 = p3.add_run(f"Назначение: {description}")
    r3.font.name   = "Times New Roman"
    r3.font.size   = Pt(9)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(60, 60, 60)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(2)

    sep2 = doc.add_paragraph("─" * 85)
    sep2.runs[0].font.name  = "Courier New"
    sep2.runs[0].font.size  = Pt(8)
    sep2.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    sep2.paragraph_format.space_before = Pt(0)
    sep2.paragraph_format.space_after  = Pt(2)

def add_code(doc, code_str):
    lines = code_str.splitlines()
    for i, line in enumerate(lines, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.2)
        num = p.add_run(f"{i:4d}  ")
        num.font.name  = "Courier New"
        num.font.size  = Pt(8.5)
        num.font.color.rgb = RGBColor(160, 160, 160)
        code = p.add_run(line)
        code.font.name  = "Courier New"
        code.font.size  = Pt(8.5)
        code.font.color.rgb = RGBColor(20, 20, 20)

# ─── main ────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Поля страницы A4: левое 3 см, правое 1.5 см, верх/низ 2 см
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    # ── Титульная страница ────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ЛИСТИНГ ПРОГРАММНОГО КОДА")
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.bold = True

    doc.add_paragraph()

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Дипломная работа")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Мобильное приложение «Farme»")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(14)
    r3.font.italic = True

    doc.add_paragraph()
    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("Сельскохозяйственный маркетплейс для Кыргызстана")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        "Платформа: Android (Java, API 26+)",
        "Backend: Firebase Realtime Database + Firebase Auth + FCM",
        "Архитектура: Activity/Fragment + RecyclerView + ViewPager2",
        "Количество исходных файлов: 49",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(line)
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(11)

    add_page_break(doc)

    # ── Оглавление ────────────────────────────────────────────
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = toc_title.add_run("СОДЕРЖАНИЕ")
    rt.font.name = "Times New Roman"
    rt.font.size = Pt(14)
    rt.font.bold = True

    doc.add_paragraph()

    current_section = None
    lnum = 0
    for entry in SECTIONS:
        sec, sec_title, fname, _, descr = entry
        lnum += 1
        lnum_str = f"{lnum:2d}."

        if sec is not None:
            current_section = sec
            # Заголовок раздела в оглавлении
            p = doc.add_paragraph()
            rb = p.add_run(f"Раздел {sec}. {sec_title}")
            rb.font.name = "Times New Roman"
            rb.font.size = Pt(11)
            rb.font.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(1)

        p = doc.add_paragraph()
        rr = p.add_run(f"    {lnum_str} {fname}")
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    add_page_break(doc)

    # ── Листинги ──────────────────────────────────────────────
    current_section = None
    lnum = 0
    missing = 0

    for entry in SECTIONS:
        sec, sec_title, fname, rel_path, descr = entry
        lnum += 1

        if sec is not None:
            add_section_heading(doc, sec, sec_title)

        full_path = os.path.join(SRC_BASE, rel_path)
        if not os.path.exists(full_path):
            print(f"  [ПРОПУЩЕН] {fname} — файл не найден: {full_path}")
            missing += 1
            continue

        with open(full_path, "r", encoding="utf-8") as f:
            code = f.read()

        add_file_header(doc, str(lnum), fname, f"app/src/main/java/com/example/farme/{rel_path}", descr)
        add_code(doc, code)

    doc.save(OUT_PATH)
    total = len(SECTIONS) - missing
    print(f"Готово: {OUT_PATH}")
    print(f"Файлов в листинге: {total} из {len(SECTIONS)}")
    if missing:
        print(f"Пропущено (не найдено): {missing}")

if __name__ == "__main__":
    build()
